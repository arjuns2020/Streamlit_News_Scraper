import os
import pandas as pd
import re
import textwrap
from gnews import GNews
from datetime import datetime, timedelta
from langid import classify
from joblib import Parallel, delayed
import streamlit as st
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.cluster import KMeans
from sklearn.metrics import pairwise_distances_argmin_min
from wordcloud import WordCloud, STOPWORDS
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from io import BytesIO
import nltk
from io import StringIO

nltk.download('punkt')
# Function to check if the article is in English
def is_english(text):
    lang, _ = classify(text)
    return lang == 'en'

# Function to get full article using GNews and perform NLP
def get_full_article_with_nlp(url):
    article = GNews().get_full_article(url)
    try:
        article.download()
        article.parse()
        article.nlp()
    except Exception as e:
        print(f"Error performing NLP on the article: {e}")
        return None
    return {
        'title': article.title,
        'text': article.text,
        'authors': article.authors,
        'summary': article.summary,
        'keywords': article.keywords,
    }

def process_article(article):
    source = article.get('source', 'N/A')
    if source == 'N/A':
        url = article.get('url', '')
        source = url.split('/')[2] if url else 'N/A'
    if not is_english(article.get('title', '')):
        return None
    title_parts = re.split('... - | - ', article.get('title', 'N/A'))
    article_source = title_parts[1].strip() if len(title_parts) > 1 else 'N/A'
    result = {'Title1': article.get('title', 'N/A'), 'Source': source, 'Article_Source': article_source, 'URL': article.get('url', 'N/A')}
    full_article = get_full_article_with_nlp(article.get('url'))
    if full_article:
        result.update(full_article)
    return result


def fetch_and_process_news(query, start_date, end_date):
    GNews.language = 'en'
    google_news = GNews(country='US', max_results=100)
    google_news.start_date = start_date
    google_news.end_date = end_date
    news_results = google_news.get_news(query)
    results_list = Parallel(n_jobs=-1)(delayed(process_article)(article) for article in news_results)
    results_list = [result for result in results_list if result is not None]
    df = pd.DataFrame(results_list)
    df = df.dropna(subset=['Title1', 'text'])
    df = df[['Article_Source', 'title', 'text', 'summary', 'keywords', 'URL']]
    return df

def cluster_and_generate_wordclouds(df):
    df['processed_summary'] = df['summary'].apply(lambda x: re.sub('[^A-Za-z]+', ' ', str(x)).lower())
    count_vectorizer = CountVectorizer(stop_words='english', max_df=0.9)
    count_matrix = count_vectorizer.fit_transform(df['processed_summary'])
    num_clusters = 5
    km = KMeans(n_clusters=num_clusters, n_init=10)
    km.fit(count_matrix)
    clusters = km.labels_.tolist()
    texts_per_cluster = [[df.iloc[i]['summary'] for i in range(len(clusters)) if clusters[i] == j] for j in range(num_clusters)]
    for i, texts in enumerate(texts_per_cluster):
        if not texts:
            texts_per_cluster[i] = ["No articles assigned to this cluster."]
    summaries = [texts_per_cluster[j][0] for j in range(num_clusters)]
    doc = Document()
    doc.add_heading('Cluster Summaries and Word Clouds', 0)
    for i, summary in enumerate(summaries):
        with st.container():
            expander = st.expander(f"Cluster {i+1} Summary and Word Cloud", expanded=True)
            with expander:
                formatted_summary = "\n".join(textwrap.wrap(summary, width=850))
                st.write(formatted_summary)
                doc.add_heading(f'Cluster {i+1} Summary:', level=1)
                doc.add_paragraph(formatted_summary)
                wordcloud = WordCloud(width=800, height=400, background_color='white', stopwords=STOPWORDS, min_font_size=10).generate(str(summary))
                plt.figure(figsize=(8, 8), facecolor=None)
                plt.imshow(wordcloud)
                plt.axis("off")
                plt.tight_layout(pad=0)
                st.pyplot(plt)
                wordcloud_image_stream = BytesIO()
                plt.savefig(wordcloud_image_stream, format='png')
                plt.close()
                wordcloud_image_stream.seek(0)
                doc.add_picture(wordcloud_image_stream, width=Inches(6))
                doc.add_paragraph()
    output_doc = BytesIO()
    doc.save(output_doc)
    output_doc.seek(0)
    return output_doc


def setup_streamlit_ui():
    st.title("News Scraper and Analyzer")
    with st.form(key='search_form'):
        #query = st.text_input('Search Query', value=st.session_state.query if 'query' in st.session_state else "")
        query = st.text_input("Enter Keywords (separated by space)", 'aws cloud datacenter')
        start_date = st.date_input('Start Date', value=st.session_state.start_date if 'start_date' in st.session_state else datetime.today() - timedelta(days=7))
        end_date = st.date_input('End Date', value=st.session_state.end_date if 'end_date' in st.session_state else datetime.today())
        search_button = st.form_submit_button(label='Search')
        
    return query, start_date, end_date, search_button

def main():
    query, start_date, end_date, search_button = setup_streamlit_ui()
    
    if search_button:
        st.session_state.query = query
        st.session_state.start_date = start_date
        st.session_state.end_date = end_date

        with st.spinner("Please wait while scraping the web..."):
            st.session_state.df = fetch_and_process_news(query, start_date, end_date)  # Placeholder, replace with your function
            
            st.session_state.output_doc = cluster_and_generate_wordclouds(st.session_state.df)  # Placeholder, replace with your function

    if 'df' in st.session_state and not st.session_state.df.empty:
        st.subheader("Results")
        st.dataframe(st.session_state.df)
        
        # Convert DataFrame to CSV
        towrite = BytesIO()
        st.session_state.df.to_csv(towrite, index=False, header=True)
        towrite.seek(0)

        # Create a link for downloading the CSV file
        st.download_button(
            label="Download CSV",
            data=towrite,
            file_name="search_results.csv",
            mime='text/csv',
    )
    
    if 'output_doc' in st.session_state:
        file_name = f"{query.replace(' ', '_')}_cluster_summary.docx"
        st.sidebar.download_button(
            label="Download Cluster Summary",
            data=st.session_state.output_doc.getvalue(),
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="cluster_summary_docx"
        )

if __name__ == "__main__":
    main()
